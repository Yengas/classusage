from lessonparser import LessonParser
from lesson import Lesson, LessonTime
from classroom import ClassRoom
from docx import Document
import re

class TableMeta:
    day_names = ["Pazartesi", "Salı", "Çarşamba", "Perşembe", "Cuma"]

    def __init__(self, row_index, days):
        self.row_index = row_index
        self.days = days

class StandartDocParser:
    time_pattern = re.compile("^([0-9]{2})\\.([0-9]{2})-([0-9]{2})\\.([0-9]{2})$")
    lesson_patterns = {
        "title": re.compile("([a-zA-ZçıüğöşİĞÜÖŞÇ]{3,})\s*([0-9]{3,})([ a-zA-Z0-9çıüğöşİĞÜÖŞÇ\.?!:;-—\[\]’“”\/,]+)", re.M),
        "room": re.compile("([A-Z])\\/([A-Z0-9])-([0-9])"),
        "teacher": re.compile("^([\\s\\.a-zA-ZçıüğöşİĞÜÖŞÇ]+)(\\(|$)", re.M)
    }

    '''Parses a block of text to create a lesson object.
    Parameters
    ----------
    text : string
        Block of text representing a cell of a table row.
    '''
    def parseLesson(self, cell, time):
        if not cell.text.strip():
            return None

        result = []
        texts = cell.text.strip().split("\n\n")

        for text in texts:
            title = StandartDocParser.lesson_patterns["title"].search(text)
            teacher = StandartDocParser.lesson_patterns["teacher"].search(text)
            room = self.getRoom(text)
            optional = self.getOptional(text)

            result.append(
                Lesson(
                    title.group(3).strip(),
                    title.group(1).strip(),
                    title.group(2).strip(),
                    teacher.group(1).strip(),
                    room,
                    optional,
                    time
                )
            )

        return result

    def getRoom(self, text):
        for predefined in ClassRoom.predefined:
            if predefined in text.lower():
                return ClassRoom.predefined[predefined]
        room = StandartDocParser.lesson_patterns["room"].search(text)
        if not room:
            return ClassRoom(None, None, None, "Undefined")
        return ClassRoom(room.group(2), room.group(1), room.group(3))

    def getOptional(self, text):
        # Department wide selective
        if "(Bölüm Seç" in text:
            return 1
        # College wide selective
        if "(Üniversite/" in text:
            return 2
        # Required
        return 0

    def getMeta(self, document):
        table = document.tables[0]
        start = 0
        while start < len(table.rows):
            if table.rows[start].cells[1].text.strip() in TableMeta.day_names:
                days = list(map(lambda cell: TableMeta.day_names.index(cell.text.strip()), table.rows[start].cells[1:]))
                return TableMeta(start, days);
            start += 1
        return None

    def getLessons(self, row, meta):
        # Return empty if there are no cells.
        if len(row.cells) <= 0:
            return []

        time = self.parseTime(row.cells[0])
        if time is None:
            return []

        result = []
        # Parse lesson object for other cells
        for idx, cell in enumerate(row.cells[1:]):
            lesson = self.parseLesson(cell, LessonTime(meta.days[idx], time[0], time[1]))

            if lesson:
                result.extend(lesson);

        return result

    def parseTime(self, cell):
        # Check if the first row is the time row.
        match = StandartDocParser.time_pattern.match(cell.text.strip())
        if match is None:
            return None

        return [
            int(match.group(1)) * 60 + int(match.group(2)),
            int(match.group(3)) * 60 + int(match.group(4))
        ]

class DocLessonParser(LessonParser):
    def getSupportedExtensions(self):
        return ["docx"]

    def canParse(self, path):
        return path.endswith('docx')

    def parse(self, path):
        document = Document(path)
        parser = StandartDocParser()
        results = []

        # Parse metadata and initialize the parsing.
        # meta contains info such as days.
        meta = parser.getMeta(document);
        if meta is None:
            return None

        # Parse each row and get lessons
        for row in document.tables[0].rows[meta.row_index + 1:]:
            lessons = parser.getLessons(row, meta)
            results = Lesson.mergeLessons(results, lessons);

        return results